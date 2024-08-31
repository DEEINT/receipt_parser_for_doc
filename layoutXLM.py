from pathlib import Path
import numpy as np
import pandas as pd
import torch
from PIL import Image, ImageDraw
import json
import os
import datasets
from transformers import LayoutLMv2ForTokenClassification
from transformers import LayoutXLMProcessor
from docx import Document
from docx.shared import Pt
from docxtpl import DocxTemplate
import cv2
import datetime

labels = ['O', 'B-COMPANY', 'I-COMPANY', 'B-DATE', 'I-DATE', 'B-ADDRESS',
          'I-ADDRESS', 'B-TOTAL', 'I-TOTAL', 'B-NUMBER', 'I-NUMBER']

id2label = {v: k for v, k in enumerate(labels)}
label2id = {k: v for v, k in enumerate(labels)}

def initialize_model():
    model = LayoutLMv2ForTokenClassification.from_pretrained("D:\диплом\сайт_для_диплома\layoutXLM_rus_trained",
                                                             id2label=id2label,
                                                             label2id=label2id,
                                                             num_labels=len(labels))
    return model

def initialize_processor():
    processor = LayoutXLMProcessor.from_pretrained("microsoft/layoutxlm-base", apply_ocr=True, ocr_lang='rus')
    return processor

def unnormalize_1000_box(bbox, width, height):
     return [
         int(width * (bbox[0] / 1000)),
         int(height * (bbox[1] / 1000)),
         int(width * (bbox[2] / 1000)),
         int(height * (bbox[3] / 1000)),
     ]

def image_preporation(im, processor, model):
    encoding = processor(im, return_offsets_mapping=True, return_tensors="pt",
                     truncation=True, padding='max_length', max_length=512)
    offset_mapping = encoding.pop('offset_mapping')
    outputs = model(**encoding)
    return outputs, encoding, offset_mapping

def get_informative_tokens_text(encoding, offset_mapping, outputs, id2label, processor):
  input_ids = encoding['input_ids'].squeeze().tolist()
  is_subword = np.array(offset_mapping.squeeze().tolist())[:,0] != 0
  predictions = outputs.logits.argmax(-1).squeeze().tolist()
  true_predictions = [id2label[pred] for idx, pred in enumerate(predictions) if not is_subword[idx]]
  full_words = []
  word_idx = []
  for i in range(len(is_subword) - 1):
    word_idx.append(i)
    if is_subword[i + 1]:
      continue
    else:
      full_words.append(''.join([processor.tokenizer.decode(input_ids[id]) for id in word_idx]))
      word_idx = []
  # adding 'end' token
  full_words.append('</s>')
  inf_pred_word_tuples = []
  for i in range(len(true_predictions)):
    if true_predictions[i] != id2label[0]:
      if (full_words[i] != '<pad>'):
       inf_pred_word_tuples.append((true_predictions[i], full_words[i]))
  return inf_pred_word_tuples

def qr_code_data_extraction(np_image):
    detector = cv2.QRCodeDetector()

    data, bbox, straight_qrcode = detector.detectAndDecode(np_image)
    if bbox is not None:
        total = data[(data.find('s=') + 2):(data.find('&fn'))]
        dateD = data[(data.find('t=') + 8):(data.find('T'))]
        dateM = data[(data.find('t=') + 6):(data.find(dateD))]
        dateY = data[(data.find('t=') + 2):(data.find(dateM))]
        date = dateD + '.' + dateM + '.' + dateY
        number = data[(data.find('i=') + 2):(data.find('&fp'))]
        if (total.isdigit() != True):
          total, date, number = None, None, None
    else:
        total, date, number = None, None, None

    return total, date, number

def im_analiz(encoding, offset_mapping, outputs, processor, im):
    nerv = get_informative_tokens_text(encoding, offset_mapping, outputs, id2label, processor)
    np_im = np.array(im)
    sum = 0.0
    return nerv, np_im, sum

def pars_text(nerv, np_im):

  adress_total = ''
  total_total = ''
  date_total = ''
  company_total = ''
  number_total = ''
  total = ''
  date = ''
  company = ''
  adress = ''
  number = ''

  for i in range(len(nerv)):
    if 'ADDRESS' in nerv[i][0]:
        if 'B' in nerv[i][0]:
          adress = nerv[i][1]
        if 'I' in nerv[i][0]:
          adress = adress + ' ' + nerv[i][1]
    if 'DATE' in nerv[i][0]:
        if 'B' in nerv[i][0]:
          date = nerv[i][1]
        if 'I' in nerv[i][0]:
          date = date + ' ' + nerv[i][1]
    if 'COMPANY' in nerv[i][0]:
        if 'B' in nerv[i][0]:
          company = nerv[i][1]
        if 'I' in nerv[i][0]:
          company = company + ' ' + nerv[i][1]
    if 'TOTAL' in nerv[i][0]:
        if 'B' in nerv[i][0]:
          total = nerv[i][1]
        if 'I' in nerv[i][0]:
          total = total + nerv[i][1]
    if 'NUMBER' in nerv[i][0]:
        if 'B' in nerv[i][0]:
          number = nerv[i][1]
        if 'I' in nerv[i][0]:
          number = number + nerv[i][1]
    if len(total) > len(total_total):
      total_total = total
    if len(date) > len(date_total):
      date_total = date
    if len(company) > len(company_total):
      company_total = company
    if len(adress) > len(adress_total):
      adress_total = adress
    if len(number) > len(number_total):
      number_total = number

  totalqr, dateqr, numberqr = qr_code_data_extraction(np_im)
  if(dateqr != '..' and dateqr != None):
    total_total = totalqr
    date_total = dateqr
    number_total = numberqr

  ttt = []
  ttt.append(adress_total)
  ttt.append(total_total.replace('=', ''))
  ttt.append(date_total.replace(':', ''))
  ttt.append(company_total)
  ttt.append(number_total.replace('<pad>', ''))
  if ttt[1] == '':
    ttt[1] = '0.00'
  pr_sum = float(ttt[1])
  return ttt, pr_sum

def preparation_report(nerv, np_im, sum, doc_num, FIO, phone, post):
    date_now = datetime.datetime.now()
    gg = date_now.year
    mm = date_now.month
    dd = date_now.day
    doc = DocxTemplate("./down/example.docx")
    context = { 'FIO' : FIO, 'phone': phone, 'post':post, 'day':dd, 'month': mm, 'ear':gg}
    doc.render(context)
    doc.save('./down/replased_not_full.docx')
    ttt, pr_sum = pars_text(nerv, np_im)
    if (doc_num == 0):
        doc = Document('./down/replased_not_full.docx')
    if (doc_num == 1):
       doc = Document('./down/replased.docx')
    table = doc.tables[1]
    i = -1
    row_num = 0
    f = 0
    for row in table.rows:
        i = i + 1
        if (row.cells[0].text == '' and f == 0):
            f = 1
            row_num = i
    j = 0
    for cell in table.rows[row_num].cells:
        if (j == 0 and cell.text == ''):
            cell.text = str(row_num-3)+'.'
            rc = cell.paragraphs[0].runs[0]
            rc.font.size = Pt(9)
        if (j == 1 and cell.text == ''):
            cell.text = ttt[2]
            rc = cell.paragraphs[0].runs[0]
            rc.font.size = Pt(9)
        if (j == 2 and cell.text == ''):
            cell.text = ttt[4][0:len(ttt[4])//2]
            rc = cell.paragraphs[0].runs[0]
            rc.font.size = Pt(9)
        if (j == 3 and cell.text == ''):
            cell.text = ttt[3]
            rc = cell.paragraphs[0].runs[0]
            rc.font.size = Pt(9)
        if (j == 4 and cell.text == ''):
            cell.text = ttt[1]
            rc = cell.paragraphs[0].runs[0]
            rc.font.size = Pt(9)
        j = j + 1

    sum = pr_sum + sum
    table.rows[27].cells[4].text = str(sum)
    doc.save('./down/replased.docx')

def work_with_im(im, processor, model, doc_num, FIO, phone, post):
   outputs, encoding, offset_mapping = image_preporation(im, processor, model)
   inf = get_informative_tokens_text(encoding, offset_mapping, outputs, id2label, processor)
   ner, np_im, sum = im_analiz(encoding, offset_mapping, outputs, processor, im)
   preparation_report(ner, np_im, sum, doc_num, FIO, phone, post)

